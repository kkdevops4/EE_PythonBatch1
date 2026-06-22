'''
pip install kaleido
pip install python-docx
'''

import pandas as pd
from analytics import *
from datetime import datetime
from docx import Document
from docx.shared import Inches


def generate_report(df):
    doc = Document()
    doc.add_heading("EV Battery Analysis Report", level = 1)

    generated_time = datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    doc.add_paragraph(f"Generated On: {generated_time}")

    doc.add_heading("Dataset Information", level = 2)
    doc.add_paragraph(f"Number of Records: {record_count}")
    doc.add_paragraph(f"Start Time: {start_time}")
    doc.add_paragraph(f"End Time: {end_time}")

    soc = soc_summary(df)
    soh = soh_summary(df)

    doc.add_heading("Battery Health Summary", level = 2)
    latest_soc = soc["Latest SoC"]
    min_soc = soc["Minimum SoC"]
    max_soc = soc["Maximum SoC"]
    latest_soh = soh["Latest SoH"]

    doc.add_paragraph(f"Latest SoC: {latest_soc}%")
    doc.add_paragraph(f"Minimum SoC: {min_soc}%")
    doc.add_paragraph(f"Maximum SoC: {max_soc}%")
    doc.add_paragraph(f"Latest SoH: {latest_soh}%")

    if latest_soh >= 95:
        health_status = "Excellent"
    elif latest_soh >= 90:
        health_status = "Good"
    elif latest_soh >= 80:
        health_status = "Moderate Degradation"
    else: 
        health_status = "Attention Required "

    doc.add_paragraph(f"Battery Health Status: {health_status}")

    doc.add_heading("Charging Analysis",level=2)
    charging_rows = (df["state"] == "Charging").sum()
    charging_hours = round((charging_rows * 10) / 3600, 2)
    doc.add_paragraph(f"Total Charging Duration: {charging_hours} hours")

    if charging_hours == 0:
        charging_status = "No charging activity detected"
    elif charging_hours < 2:
        charging_status = "Light charging activity"
    elif charging_hours < 5:
        charging_status = "Moderate charging activity"
    else:
        charging_status = "High charging activity"  
    doc.add_paragraph(f"Charging Activity Status: {charging_status}")

    doc.add_heading("State of Charge Trend", level = 2)
    doc.add_picture("output_data/soc_chart.png", width=Inches(6))

    doc.save("output_data/battery_report.docx")

df = load_data("processed_data/processed_battery_data.xlsx")

record_count = len(df)
start_time = df["timestamp"].min()
end_time = df["timestamp"].max()

generate_report(df)

print("word document created sucessfully")