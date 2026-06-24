# =====================================================
# INSTALLS
# =====================================================
'''
pip install python-docx
'''

# =====================================================
# IMPORTS
# =====================================================
import pandas as pd
from modules.analytics import *
from modules.charts import *
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx2pdf import convert


def generate_report(df):
# =====================================================
# CALCULATIONS
# =====================================================
    record_count = len(df)
    start_time = df["timestamp"].min()
    end_time = df["timestamp"].max()

    soc = soc_summary(df)
    soh = soh_summary(df)
    voltage = voltage_summary(df)

    latest_soc = soc["Latest SoC"]
    min_soc = soc["Minimum SoC"]
    max_soc = soc["Maximum SoC"]
    latest_soh = soh["Latest SoH"]

    latest_voltage = voltage["Latest Voltage"]
    max_temp = round(df["battery_temp_c"].max(), 2)

#HEALTH STATUS
    health_status = battery_health_status(latest_soh)

#OVERALL STATUS
    overall_status = overall_condition(latest_soh,max_temp)

#CHARGING HOURS
    charging = charging_summary(df)

    charging_hours = charging["Charging Hours"]
    charging_status = charging["Charging Status"]

# =====================================================
# DOCUMENT CREATION
# =====================================================
    doc = Document()
    doc.add_heading("EV Battery Analysis Report", level = 1)

    generated_time = datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    doc.add_paragraph(f"Generated On: {generated_time}")

# =====================================================
# EXECUTIVE SUMMARY
# =====================================================
    doc.add_heading("Executive Summary", level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "Metric"
    header[1].text = "Value"

    row = table.add_row().cells
    row[0].text = "Latest SoC"
    row[1].text = f"{latest_soc}%"

    row = table.add_row().cells
    row[0].text = "Latest SoH"
    row[1].text = f"{latest_soh}%"

    row = table.add_row().cells
    row[0].text = "Latest Voltage"
    row[1].text = f"{latest_voltage} V"

    row = table.add_row().cells
    row[0].text = "Maximum Temperature"
    row[1].text = f"{max_temp} °C"

    row = table.add_row().cells
    row[0].text = "Battery Health"
    row[1].text = health_status

    row = table.add_row().cells
    row[0].text = "Charging Duration"
    row[1].text = f"{charging_hours} Hours"

    row = table.add_row().cells
    row[0].text = "Overall Status"
    row[1].text = overall_status

# =====================================================
# DATASET INFORMATION
# =====================================================
    doc.add_heading("Dataset Information", level = 2)
    doc.add_paragraph(f"Number of Records: {record_count}")
    doc.add_paragraph(f"Start Time: {start_time}")
    doc.add_paragraph(f"End Time: {end_time}")

# =====================================================
# BATTERY HEALTH SUMMARY
# =====================================================
    doc.add_heading("Battery Health Summary", level = 2)


    doc.add_paragraph(f"Latest SoC: {latest_soc}%")
    doc.add_paragraph(f"Minimum SoC: {min_soc}%")
    doc.add_paragraph(f"Maximum SoC: {max_soc}%")
    doc.add_paragraph(f"Latest SoH: {latest_soh}%")

    doc.add_paragraph(f"Battery Health Status: {health_status}")

# =====================================================
# CHARGING ANALYSIS
# =====================================================
    doc.add_heading("Charging Analysis",level=2)

    doc.add_paragraph(f"Total Charging Duration: {charging_hours} hours")

    if charging_hours == 0:
        charging_status = "No charging activity detected"
    elif charging_hours < 2:
        charging_status = "Light charging activity"
    elif charging_hours < 5:
        charging_status = "Moderate charging activity"
    else:
        charging_status = "High charging activity"  
    doc.add_paragraph(f"Charging Activity Status: {charging_status}")

# =====================================================
# STATE OF CHARGE CHART
# =====================================================
    doc.add_heading("State of Charge Trend", level = 2)
    doc.add_picture("output_data/soc_chart.png", width=Inches(6))
    doc.add_paragraph("Figure 1: State of Charge variation over time.")

# =====================================================
# VOLTAGE CHART
# =====================================================
    doc.add_heading("Voltage Analysis", level=2)
    doc.add_picture("output_data/voltage_chart.png",width=Inches(6))
    doc.add_paragraph("Figure 2: Battery pack voltage variation over time.")

    min_voltage = round(df["battery_voltage_v"].min(), 2)
    max_voltage = round(df["battery_voltage_v"].max(), 2)

    doc.add_paragraph(f"Voltage ranged from {min_voltage} V to {max_voltage} V.")

# =====================================================
# CURRENT CHART
# =====================================================
    doc.add_heading("Current Analysis", level=2)
    doc.add_picture("output_data/current_chart.png",width=Inches(6))
    doc.add_paragraph("Figure 3: Battery pack current variation over time.")

    max_current = round(df["battery_current_a"].max(), 2)
    min_current = round(df["battery_current_a"].min(), 2)

    doc.add_paragraph(f"Maximum discharge current: {max_current} A")

    doc.add_paragraph(f"Maximum charging current: {min_current} A")

# =====================================================
# TEMPERATURE CHART
# =====================================================
    doc.add_heading("Temperature Analysis", level=2)
    doc.add_picture("output_data/temp_chart.png",width=Inches(6))
    doc.add_paragraph("Figure 4: Battery pack temperature variation over time.")

    doc.add_paragraph(f"Maximum battery temperature recorded was {max_temp} °C.")

    if max_temp > 45:
        temp_status = "Warning"
    else:
        temp_status = "Normal"
    doc.add_paragraph(f"Temperature Status: {temp_status}")

    doc.add_heading("Alerts and Warnings", level=2)

# =====================================================
# ALERTS
# =====================================================
    alerts = []
    if max_temp > 45:
        alerts.append("Temperature exceeded safe operating threshold.")
    if latest_soh < 80:
        alerts.append("Battery health degradation detected.")
    if len(alerts) == 0:
        doc.add_paragraph("No critical alerts detected.")
    else:
        for alert in alerts:
            doc.add_paragraph(alert)


# =====================================================
# CONCLUSION SUMMARY
# =====================================================
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        f"""
    The battery operated with a maximum temperature of {max_temp}°C,
    maintained a State of Health of {latest_soh}%,
    and showed {charging_status.lower()}.
    Overall battery condition was assessed as {overall_status}.
    """
    )

# =====================================================
# SAVE DOCUMENT
# =====================================================
    docx_file = "output_data/battery_report.docx"
    pdf_file = "output_data/battery_report.pdf"
    
    try:
        doc.save(docx_file)
        print("Word document created successfully")
    except Exception as e:
        print("Word document creation failed")
        print(e)

    try:
        convert(docx_file, pdf_file)
        print("PDF report created successfully")
    except Exception as e:
        print("PDF conversion failed")
        print(e)


def create_report():
# =====================================================
# LOAD DATA
# =====================================================
    df = load_data("processed_data/processed_battery_data.xlsx")

# =====================================================
# CREATE CHARTS 
# =====================================================
    create_soc_chart(df)
    create_voltage_chart(df)
    create_current_chart(df)
    create_temp_chart(df)

# =====================================================
# GENERATE REPORT
# =====================================================
    generate_report(df)
    print("Report generated successfully")
 
    
if __name__ == "__main__":
    create_report()