import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from datetime import datetime
from abc import ABC, abstractmethod


# ==========================================================
# FILE HANDLER (UPDATED FOR 2 SHEETS)
# ==========================================================

class FileHandler:

    @staticmethod
    def read_excel(file_path):
        try:
            if not file_path.endswith(".xlsx"):
                raise ValueError("Only .xlsx files allowed")

            service_card = pd.read_excel(file_path, sheet_name="Service Card")
            vehicle_data = pd.read_excel(file_path, sheet_name="Vehicle Data")

            if service_card.empty or vehicle_data.empty:
                raise ValueError("One or more sheets are empty")

            return service_card, vehicle_data

        except Exception as e:
            print("File Error:", e)
            return None, None


# ==========================================================
# VEHICLE CLASS
# ==========================================================

class Vehicle:
    def __init__(self, service_card, vehicle_data):
        self.service_card = service_card
        self.vehicle_data = vehicle_data


# ==========================================================
# COMPARATOR (NEW LOGIC)
# ==========================================================

class VehicleComparator:

    def compare(self, standard_df, current_df):

        results = {}

        merged = pd.merge(
            standard_df,
            current_df,
            on="Parameter",
            how="inner"
        )

        for _, row in merged.iterrows():

            param = row["Parameter"]
            standard_value = row["Standard Condition"]
            current_value = row["Today's Condition"]

            if standard_value == 0:
                deviation = 0
            else:
                deviation = abs(current_value - standard_value) / standard_value * 100

            results[param] = round(deviation, 2)

        return results


# ==========================================================
# PREDICTOR
# ==========================================================

class BasePredictor(ABC):
    @abstractmethod
    def predict(self, deviation):
        pass


class WearPredictor(BasePredictor):

    def predict(self, deviation):
        if deviation <= 10:
            return "Good Condition"
        elif deviation <= 25:
            return "Moderate Wear"
        else:
            return "Critical Condition"


# ==========================================================
# ALERT SYSTEM
# ==========================================================

class AlertGenerator:

    def generate(self, parameter, status):

        if status != "Critical Condition":
            return "No Alert"

        alerts = {
            "Brake Condition": "Brake failure risk",
            "Engine Temperature": "Engine overheating",
            "Battery Health": "Battery failure risk",
            "Oil Level": "Oil shortage",
            "Tire Pressure": "Tyre damage risk",
            "Coolant Level": "Cooling system issue",
            "Fuel Level": "Fuel leakage or shortage"
        }

        return alerts.get(parameter, "Immediate inspection required")


# ==========================================================
# VISUALIZER
# ==========================================================

class Visualizer:

    def bar(self, deviations):
        plt.figure(figsize=(12, 5))
        plt.bar(deviations.keys(), deviations.values(), color="orange")
        plt.xticks(rotation=45)
        plt.title("Vehicle Parameter Deviation")
        plt.tight_layout()
        plt.savefig("deviation.png")
        plt.close()

    def pie(self, status_count):
        plt.figure(figsize=(6, 6))
        plt.pie(status_count.values(), labels=status_count.keys(), autopct="%1.1f%%")
        plt.title("Vehicle Health Distribution")
        plt.savefig("health.png")
        plt.close()


# ==========================================================
# REPORT GENERATOR (SERVICE CARD ON TOP)
# ==========================================================

class ReportGenerator:

    def create(self, service_card, report_rows, avg_dev, overall_status):

        doc = Document()

        # ================= SERVICE CARD (TOP) =================
        doc.add_heading("SERVICE CARD", level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"

        for col in service_card.columns:
            row = table.add_row().cells
            row[0].text = str(col)
            row[1].text = str(service_card.iloc[0][col])

        doc.add_page_break()

        # ================= SUMMARY =================
        doc.add_heading("Vehicle Health Summary", level=1)
        doc.add_paragraph(f"Average Deviation: {avg_dev}%")
        doc.add_paragraph(f"Overall Status: {overall_status}")

        # ================= ANALYSIS TABLE =================
        doc.add_heading("Detailed Analysis", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Parameter"
        hdr[1].text = "Deviation"
        hdr[2].text = "Status"
        hdr[3].text = "Alert"

        for r in report_rows:
            row = table.add_row().cells
            row[0].text = r["Parameter"]
            row[1].text = str(r["Deviation"])
            row[2].text = r["Status"]
            row[3].text = r["Alert"]

        # ================= IMAGES =================
        doc.add_page_break()
        doc.add_heading("Graphs", level=1)

        doc.add_picture("deviation.png", width=Inches(6))
        doc.add_picture("health.png", width=Inches(5))

        file_name = "Vehicle_Report.docx"
        doc.save(file_name)

        print("DOCX Generated:", file_name)

        try:
            convert(file_name)
            print("PDF Generated: Vehicle_Report.pdf")
        except:
            print("PDF conversion failed")


# ==========================================================
# MAIN SYSTEM
# ==========================================================

def run_system():

    standard_file = input("Enter Standard File: ")
    current_file = input("Enter Current File: ")

    std_service, std_data = FileHandler.read_excel(standard_file)
    cur_service, cur_data = FileHandler.read_excel(current_file)

    if std_service is None or std_data is None or cur_service is None or cur_data is None:
        print("Failed to load files. Please check your input files.")
        return

    vehicle_std = Vehicle(std_service, std_data)
    vehicle_cur = Vehicle(cur_service, cur_data)

    comparator = VehicleComparator()
    deviations = comparator.compare(
        vehicle_std.vehicle_data,
        vehicle_cur.vehicle_data
    )

    if not deviations:
        print("No matching parameters found between standard and current data.")
        return

    predictor = WearPredictor()
    alert = AlertGenerator()

    status_count = {
        "Good Condition": 0,
        "Moderate Wear": 0,
        "Critical Condition": 0
    }

    report_rows = []

    for param, dev in deviations.items():

        status = predictor.predict(dev)
        a = alert.generate(param, status)

        status_count[status] += 1

        report_rows.append({
            "Parameter": param,
            "Deviation": dev,
            "Status": status,
            "Alert": a
        })

    avg_dev = round(sum(deviations.values()) / len(deviations), 2)

    if avg_dev <= 10:
        overall = "Good Condition"
    elif avg_dev <= 25:
        overall = "Moderate Wear"
    else:
        overall = "Critical Condition"

    viz = Visualizer()
    viz.bar(deviations)
    viz.pie(status_count)

    report = ReportGenerator()
    report.create(
        vehicle_std.service_card,
        report_rows,
        avg_dev,
        overall
    )


if __name__ == "__main__":
    run_system()