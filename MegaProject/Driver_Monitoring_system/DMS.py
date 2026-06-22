import os
import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx.shared import Inches
from docx2pdf import convert

# CREATE OUTPUT FOLDER

os.makedirs("MegaProject/Driver_Monitoring_system/Report/graphs", exist_ok=True)       # exist_ok=True  If folder already exists don't give error. if we can't use this python crash second time.

# FUNCTIONS

def get_driver_state(score):
    if score >= 85:
        return "Alert"
    elif score >= 70:
        return "Moderate"
    elif score >= 50:
        return "Slight_Fatigue"
    else:
        return "Drowsy"

def get_driver_remark(attention):

    if attention >= 85:
        return ("Driver is highly alert and maintaining good attention throughout the journey. No signs of fatigue detected.")

    elif attention >= 70:
        return ("Driver shows moderate attention levels. Short breaks are recommended during long trips.")

    elif attention >= 50:
        return ("Driver is showing signs of fatigue. Continuous driving may affect safety and concentration.")

    else:
        return ("Driver is highly drowsy and at significant risk. Immediate rest is strongly recommended.")


def generate_driver_report(row, remark):

    return f"""
Driver ID           : {row['Driver_ID']}
Driver Name         : {row['Driver_names']}

Travel Time         : {row['Total_Time_Minutes']} Minutes
Distance Covered    : {row['Total_Distance_km']:.2f} km

Average Eye Closure : {row['Avg_Eye_Closure']:.2f} %
Average Blink Rate  : {row['Avg_Blink_Rate']:.2f} blinks/min
Average Head Pitch  : {row['Avg_Head_Pitch']:.2f}°

Attention Score     : {row['Avg_Attention_Score']:.2f}
Driver State        : {row['Final_State']}
Driving Session     : {row['Time_Period']}        

Driver Assessment:
{remark}
"""

# for Driving session

def get_time_period(timestamp):      

    hour = timestamp.hour

    if 5 <= hour < 12:
        return "Morning"

    elif 12 <= hour < 17:
        return "Afternoon"

    elif 17 <= hour < 22:
        return "Night"

    else:
        return "Early Morning"


def get_time_risk(period):

    if period == "Morning":
        return "Low Risk"

    elif period == "Afternoon":
        return "Medium Risk"

    elif period == "Night":
        return "High Risk"

    else:
        return "Very High Risk"
    
# READ CSV

csv_file = r"C:\Users\ia97974\Desktop\51856\EE_PythonBatch1\MegaProject\Driver_Monitoring_system\driver_monitoring_system_1.csv"

df = pd.read_csv(csv_file)

df["Timestamp"] = pd.to_datetime(df["Timestamp"])   

df["Time_Period"] = df["Timestamp"].apply(get_time_period) 

# ATTENTION SCORE

eye_score = 100 - df["Eye_Closure_Percentage"]

blink_score = 100 - (abs(15 - df["Blink_Rate"]) * 5)

head_score = 100 - (df["Head_Pitch_Angle"] * 4)

eye_score = eye_score.clip(0, 100)
blink_score = blink_score.clip(0, 100)
head_score = head_score.clip(0, 100)

df["Attention_Score"] = (eye_score * 0.60 + blink_score * 0.25 + head_score * 0.15).round(2)


# LINE GRAPH FOR EACH DRIVER

for driver_id in df["Driver_ID"].unique():

    driver_data = df[df["Driver_ID"] == driver_id]

    driver_name = driver_data["Driver_names"].iloc[0]
    
    time_period = driver_data["Time_Period"].iloc[0] 
    risk = get_time_risk(time_period)                 

    plt.figure(figsize=(10, 6))

    plt.plot(
        driver_data["Time_Minutes"],
        driver_data["Eye_Closure_Percentage"],
        label="Eye Closure %",
        color="red",
        linewidth=2
    )

    plt.plot(
        driver_data["Time_Minutes"],
        driver_data["Blink_Rate"],
        label="Blink Rate",
        color="orange",
        linewidth=2
    )

    plt.plot(
        driver_data["Time_Minutes"],
        driver_data["Head_Pitch_Angle"],
        label="Head Pitch",
        color="purple",
        linewidth=2
    )

    plt.plot(
        driver_data["Time_Minutes"],
        driver_data["Attention_Score"],
        label="Attention Score",
        color="green",
        linewidth=3
    )

    plt.axhspan(85,100,color="green",alpha=0.15)
    plt.axhspan(70,85,color="yellow",alpha=0.15)
    plt.axhspan(50,70,color="orange",alpha=0.15)
    plt.axhspan(0,50,color="red",alpha=0.15)         
    
    plt.title(f"Driver Monitoring Trend\n {driver_name} (ID: {driver_id})\n Session: {time_period}")   

    plt.xlabel("Time (Minutes)")
    plt.ylabel("Value")

    plt.legend()
    plt.grid(True, linestyle="--")
    

    plt.text(                      
    0.02,
    0.95,
    f"Session: {time_period}",
    transform=plt.gca().transAxes,
    fontsize=11,
    bbox=dict(facecolor="white")
)

    plt.text(                      
        0.02,
        0.88,
        f"Risk: {risk}",
        transform=plt.gca().transAxes,
        fontsize=11,
        bbox=dict(facecolor="white")
    )

    plt.tight_layout()

    graph_file = (f"MegaProject/Driver_Monitoring_system/Report/graphs/Driver_{driver_id}_Attention_Trend.png")

    plt.savefig(graph_file, dpi=300, bbox_inches="tight")
    plt.close()
 
 
states = []                            # eg.- states = ["Alert", "Moderate", "Slight Fatigue"]

for score in df["Attention_Score"]:
    states.append(get_driver_state(score)) 

df["State"] = states                  # Pandas creates a new column in data frame i.e State


report = df.groupby(                     
    ["Driver_ID", "Driver_names"]
).agg({
    "Time_Period": "first",
    "Time_Minutes": "max",
    "Travel_Distance_km": "max",
    "Eye_Closure_Percentage": "mean",
    "Blink_Rate": "mean",
    "Head_Pitch_Angle": "mean",
    "Attention_Score": "mean"
}).reset_index()

avg_state = (df.groupby(["Driver_ID", "Driver_names"])["Attention_Score"].mean().reset_index())
avg_state["Driver_State"] = avg_state["Attention_Score"].apply(get_driver_state)
abc = avg_state["Driver_State"].to_list()

avg_state["Final_State"] = abc

report = report.merge(avg_state[["Driver_ID", "Driver_names", "Final_State"]],on=["Driver_ID", "Driver_names"],how="left")

report.columns = [          
    "Driver_ID",
    "Driver_names",
    "Time_Period",
    "Total_Time_Minutes",
    "Total_Distance_km",
    "Avg_Eye_Closure",
    "Avg_Blink_Rate",
    "Avg_Head_Pitch",
    "Avg_Attention_Score",
    "Final_State"
]

# CONSOLE REPORT

print("\n==============================")
print(" DRIVER WISE REPORT")
print("==============================\n")

for _, row in report.iterrows():

    attention = row["Avg_Attention_Score"]    # take the value in the column Avg_Attention_Score and store it in a variable called attention

    remark = get_driver_remark(attention)

    print(generate_driver_report(row,remark))

# INDIVIDUAL DRIVER GRAPHS

for _, row in report.iterrows():

    metrics = [
        row["Avg_Eye_Closure"],
        row["Avg_Blink_Rate"],
        row["Avg_Head_Pitch"],
        row["Avg_Attention_Score"]
        
    ]                                                      # EXAMPLE- metrics = [20, 15, 10, 85]

    labels = [
        "Eye Closure %",
        "Blink Rate",
        "Head Pitch",
        "Attention Score"
    ]

    plt.figure(figsize=(10, 6))

    bars = plt.bar(labels,metrics,color=["red","orange","purple","green"])

    plt.title(f"Driver Performance Report\n {row['Driver_names']} (ID: {row['Driver_ID']})\n Session: {row['Time_Period']}")

    plt.ylabel("Value")

# add the numerical value on top of each bar

    for bar in bars:

        plt.text(
            bar.get_x()
            + bar.get_width()/2,
            bar.get_height(),
            f"{bar.get_height():.1f}",
            ha="center",               # ha = Horizontal Alignment
            va="bottom"                # va = Vertical Alignment
        )

    plt.grid(axis="y", linestyle="--")

    plt.tight_layout()

    graph_file = (f"MegaProject/Driver_Monitoring_system/Report/graphs/Driver_{row['Driver_ID']}_Performance.png")

    plt.savefig(graph_file,dpi=300,bbox_inches="tight")                # dpi=300 High quality image, Suitable for reports and PDFs
                                                                       # bbox_inches="tight" Removes unnecessary white space around the graph.

    plt.close()

# COMPARISON GRAPH

driver_attention = report.set_index("Driver_names")["Avg_Attention_Score"]

plt.figure(figsize=(10, 6))

bars = plt.bar(driver_attention.index,driver_attention.values,color="skyblue")

for bar in bars:

    score = bar.get_height()

    if score >= 85:
        status = "Excellent"
    elif score >= 70:
        status = "Good"
    elif score >= 50:
        status = "Fatigued"
    else:
        status = "Drowsy"

    plt.text(
        bar.get_x()
        + bar.get_width()/2,
        score,
        f"{score:.1f}\n{status}",
        ha="center",
        va="bottom"
    )

plt.title("Average Attention Score by Driver")

plt.xlabel("Driver")
plt.ylabel("Attention Score")
plt.ylim(0, 100)

plt.grid(axis="y", linestyle="--")

plt.tight_layout()

comparison_graph = ("MegaProject/Driver_Monitoring_system/Report/graphs/Driver_Attention_Comparison.png")

plt.savefig(comparison_graph,dpi=300,bbox_inches="tight")
plt.close()

# WORD REPORT

doc = docx.Document()

doc.add_heading("Driver Monitoring System Report", level=0)

for _, row in report.iterrows():

    attention = row["Avg_Attention_Score"]

    remark = get_driver_remark(attention)

    doc.add_heading(f"Driver: {row['Driver_names']} (ID: {row['Driver_ID']})",level=1)

    doc.add_paragraph(generate_driver_report(row, remark))

    #############
    # BAR GRAPH
    #############

    performance_graph = (f"MegaProject/Driver_Monitoring_system/Report/graphs/Driver_{row['Driver_ID']}_Performance.png")

    if os.path.exists(performance_graph):

        doc.add_heading("Performance Summary",level=2)

        doc.add_picture(performance_graph,width=Inches(5.5))

    #############
    # LINE GRAPH
    #############

    trend_graph = (f"MegaProject/Driver_Monitoring_system/Report/graphs/Driver_{row['Driver_ID']}_Attention_Trend.png")

    if os.path.exists(trend_graph):

        doc.add_heading("Attention Score Trend",level=2)

        doc.add_picture(trend_graph,width=Inches(5.5))

    doc.add_page_break()
    
    
# OVERALL COMPARISON GRAPH

doc.add_heading("Driver Attention Score Comparison",level=1)

if os.path.exists(comparison_graph):

    doc.add_picture(comparison_graph,width=Inches(6))

# SAVE WORD FILE

doc_file = ("MegaProject/Driver_Monitoring_system/Report/Driver_Monitoring_Report.docx")
doc.save(doc_file)
print("\nWord Report Saved Successfully!")

# PDF CONVERSION

try:
    convert(doc_file)
    print("PDF Report Generated Successfully!")

except Exception as e:
    print("\nPDF Conversion Failed!")
    print(e)